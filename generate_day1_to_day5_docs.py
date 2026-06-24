#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据飞书 Wiki 文档内容生成「智汇桥」Day1-Day5 开发日志 Word 文档。

功能说明：
1. 读取 day1_content.md ~ day5_content.md（从飞书 Wiki 导出）。
2. 解析 Markdown 结构（标题、表格、列表、代码块、图片、正文）。
3. 自动下载文档中的远程图片（如 Day2 的数据库 ER 图）。
4. 为对应日期补充关键页面截图（Day2 ER 图、Day5 登录页/项目列表页）。
5. 输出 5 份格式统一的 Word 文档。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import json
import subprocess
import textwrap
from urllib.parse import urlparse
from typing import List, Dict, Any, Tuple

# ==================== 全局配置 ====================

# 工作目录
OUTPUT_DIR = r"C:\Users\luofeng\Desktop\实训"

# 截图/图片根目录
IMAGE_BASE_DIR = r"C:\Users\luofeng\Desktop\实训\docs"

# 每天的文档配置：标题、日期、内容文件、需要补充的截图
DAY_CONFIGS = [
    {
        "day": "Day1",
        "title": "智汇桥项目开发日志 - Day 1",
        "date": "15",
        "content_file": "day1_content.md",
        "screenshots": []
    },
    {
        "day": "Day2",
        "title": "智汇桥项目开发日志 - Day 2",
        "date": "16",
        "content_file": "day2_content.md",
        "screenshots": [
            {"title": "数据库ER图", "path": "day2/er-diagram-1.png", "caption": "Day2 完成的数据库ER设计图"}
        ]
    },
    {
        "day": "Day3",
        "title": "智汇桥项目开发日志 - Day 3",
        "date": "17",
        "content_file": "day3_content.md",
        "screenshots": []
    },
    {
        "day": "Day4",
        "title": "智汇桥项目开发日志 - Day 4",
        "date": "18",
        "content_file": "day4_content.md",
        "screenshots": []
    },
    {
        "day": "Day5",
        "title": "智汇桥项目开发日志 - Day 5",
        "date": "19",
        "content_file": "day5_content.md",
        "screenshots": [
            {"title": "登录页", "path": "day6/day6-01-login-page.png", "caption": "Day5 联调完成的登录页面"},
            {"title": "科研项目列表页", "path": "day6/day6-02-student-project-list.png", "caption": "Day5 完成的科研项目列表页"}
        ]
    }
]


# ==================== 字体与样式工具 ====================

def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0),
             bold=False, italic=False):
    """
    统一设置 Word 文字 run 的字体、大小、颜色、粗体、斜体。
    """
    run.font.name = font_name
    # 同时设置中文字体，避免在部分 Word 版本中中文字体回退失败
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name
    )
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """
    添加自定义标题段落。
    level=0 用于文档主标题，level=1 用于一级章节，level=2 用于二级小节。
    """
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    elif level == 2:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    else:
        set_font(run, size=12, bold=True)
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0),
                         italic=False, align=None, space_after=6,
                         first_line_indent=0.25):
    """
    添加自定义正文段落，支持对齐、首行缩进等。
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def add_rich_paragraph(doc, text, space_after=6, first_line_indent=0.25):
    """
    添加支持 Markdown 行内格式（**粗体**、*斜体*）的段落。
    将文本按粗体/斜体标记拆分为多个 run，统一渲染。
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)

    # 匹配 **粗体** 或 *斜体*，保留分隔符以便判断
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)
    return p


def add_bullet_list(doc, items):
    """添加项目符号列表。"""
    for item in items:
        add_rich_paragraph(doc, f"• {item}", space_after=4, first_line_indent=0.25)


def add_numbered_list(doc, items):
    """添加编号列表。"""
    for index, item in enumerate(items, start=1):
        add_rich_paragraph(doc, f"{index}. {item}", space_after=4, first_line_indent=0.25)


def add_table_from_data(doc, headers, rows):
    """
    根据表头和数据创建 Word 表格，并设置字体样式。
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header.strip()
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)

    # 数据行
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = value.strip()
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)
    doc.add_paragraph()
    return table


def add_code_block(doc, code_text, caption=""):
    """
    添加等宽代码块样式段落。
    """
    if caption:
        add_paragraph_custom(doc, caption, bold=True, size=10, space_after=2)
    for line in code_text.strip('\n').split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, font_name="Consolas", size=9, color=RGBColor(0x33, 0x33, 0x33))
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.2


def add_screenshot(doc, image_path, caption=""):
    """
    添加截图到文档。如果图片不存在则显示红色提示文字。
    """
    full_path = os.path.join(IMAGE_BASE_DIR, image_path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.space_after = Pt(4)
        if caption:
            add_paragraph_custom(
                doc, caption, size=9, color=RGBColor(0x66, 0x66, 0x66),
                italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
            )
    else:
        add_paragraph_custom(
            doc, f"[图片缺失：{image_path}]", bold=True,
            color=RGBColor(0xFF, 0x00, 0x00), space_after=10
        )


# ==================== Markdown 解析 ====================

def parse_table(table_lines: List[str]) -> Dict[str, Any]:
    """
    解析 Markdown 表格文本，返回 headers 和 rows。
    支持表头后第一行为分隔符（|-|-|），自动忽略。
    """
    cleaned = []
    for line in table_lines:
        # 去掉首尾的 '|'，再按 '|' 分割
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        cleaned.append(cells)

    if not cleaned:
        return None

    headers = cleaned[0]
    rows = []
    for line_cells in cleaned[1:]:
        # 跳过分隔行（全部由 '-' 组成）
        if all(re.fullmatch(r'[-:\s]+', cell) for cell in line_cells):
            continue
        # 如果列数不足，补空字符串；如果过多，截断
        row = (line_cells + [''] * len(headers))[:len(headers)]
        rows.append(row)
    return {"headers": headers, "rows": rows}


def parse_markdown_content(content: str) -> List[Dict[str, Any]]:
    """
    将 Markdown 内容解析为结构化章节列表，供 render_section 渲染。

    支持的元素：
    - 标题（# / ## / ### / ####）
    - 表格（| ... |）
    - 无序列表（- ...）
    - 有序列表（1. ...）
    - 代码块（``` ... ```）
    - 图片（![...](url) 或 <figure> 标签）
    - 正文段落（支持 **粗体**、*斜体*）
    """
    # 移除 <title> 标签
    content = re.sub(r'<title>.*?</title>', '', content, flags=re.DOTALL).strip()
    # 移除 <figure> 标签本身（保留其中的图片链接已经在 markdown 中）
    content = re.sub(r'<figure.*?</figure>', '', content, flags=re.DOTALL)

    sections = []
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_caption = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                code_caption = stripped[3:].strip()
            else:
                in_code_block = False
                sections.append({
                    "type": "code",
                    "code": "\n".join(code_lines),
                    "caption": code_caption
                })
                code_caption = ""
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # 标题
        if re.match(r'^#{1}\s+', stripped):
            sections.append({"type": "heading", "text": re.sub(r'^#\s+', '', stripped), "level": 1})
            i += 1
            continue
        if re.match(r'^#{2}\s+', stripped):
            sections.append({"type": "heading", "text": re.sub(r'^##\s+', '', stripped), "level": 1})
            i += 1
            continue
        if re.match(r'^#{3}\s+', stripped):
            sections.append({"type": "heading", "text": re.sub(r'^###\s+', '', stripped), "level": 2})
            i += 1
            continue
        if re.match(r'^#{4}\s+', stripped):
            sections.append({"type": "heading", "text": re.sub(r'^####\s+', '', stripped), "level": 3})
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            table_data = parse_table(table_lines)
            if table_data:
                sections.append({
                    "type": "table",
                    "headers": table_data["headers"],
                    "rows": table_data["rows"]
                })
            continue

        # 图片（Markdown 语法）
        # 说明：飞书导出的图片多为临时 authcode 链接，与本地已有截图重复，
        # 因此这里不自动下载，统一通过 DAY_CONFIGS 中的 screenshots 配置嵌入。
        image_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if image_match:
            i += 1
            continue

        # 无序列表
        if stripped.startswith('- '):
            items = []
            while i < len(lines):
                item_line = lines[i].strip()
                if not item_line.startswith('- '):
                    break
                # 去掉列表标记并保留后续内容
                items.append(item_line[2:])
                i += 1
            sections.append({"type": "bullet", "items": items})
            continue

        # 有序列表
        ordered_match = re.match(r'^(\d+)\.\s+', stripped)
        if ordered_match:
            items = []
            while i < len(lines):
                item_line = lines[i].strip()
                if not re.match(r'^\d+\.\s+', item_line):
                    break
                items.append(re.sub(r'^\d+\.\s+', '', item_line))
                i += 1
            sections.append({"type": "numbered", "items": items})
            continue

        # 普通段落（支持行内粗体/斜体）
        sections.append({"type": "rich_paragraph", "text": stripped})
        i += 1

    return sections


# ==================== 图片下载 ====================

def download_image(url: str, output_path: str) -> bool:
    """
    使用 lark-cli docs +media-download 或 curl/requests 下载远程图片。
    如果图片已存在则跳过。
    """
    if os.path.exists(output_path):
        return True

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 情况1：Feishu 素材 token（短 token，形如 AbCdEf123）
    # 从 <figure> 等场景提取的 token 一般为 file_token
    if re.fullmatch(r'[A-Za-z0-9_-]{10,}', url):
        cmd = [
            "lark-cli", "docs", "+media-download",
            "--token", url,
            "--output", output_path
        ]
        result = subprocess.run(" ".join(cmd), shell=True, capture_output=True, text=True)
        if result.returncode == 0 and os.path.exists(output_path):
            return True

    # 情况2：完整 URL，使用 Python requests 下载
    try:
        import requests
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200:
            with open(output_path, 'wb') as f:
                f.write(resp.content)
            return True
    except Exception as e:
        print(f"下载图片失败 [{url}]: {e}")
    return False


def process_document_images(sections: List[Dict[str, Any]], day: str) -> List[Dict[str, Any]]:
    """
    将文档中的远程图片 URL 替换为本地截图配置。
    下载成功后返回 screenshots 列表供 Word 嵌入。
    """
    screenshots = []
    image_counter = 1

    for section in sections:
        if section.get("type") == "image_url":
            url = section.get("url", "")
            # 本地保存路径
            ext = ".png"
            parsed = urlparse(url)
            if parsed.path and '.' in parsed.path:
                ext = os.path.splitext(parsed.path)[1] or ".png"
            local_path = os.path.join(IMAGE_BASE_DIR, day.lower(), f"doc-image-{image_counter:02d}{ext}")

            # 尝试下载
            if download_image(url, local_path):
                screenshots.append({
                    "title": section.get("title", f"文档图片 {image_counter}"),
                    "path": os.path.relpath(local_path, IMAGE_BASE_DIR).replace('\\', '/'),
                    "caption": section.get("title", f"文档图片 {image_counter}")
                })
                image_counter += 1
            # 从 sections 中移除，避免渲染阶段重复处理
            section["type"] = "processed"

    return screenshots


# ==================== Word 文档渲染 ====================

def render_section(doc, section):
    """
    根据章节类型渲染到 Word 文档。
    """
    section_type = section.get("type")

    if section_type == "heading":
        add_heading_custom(doc, section["text"], level=section.get("level", 2))

    elif section_type == "rich_paragraph":
        add_rich_paragraph(doc, section["text"], space_after=6, first_line_indent=0.25)

    elif section_type == "paragraph":
        add_paragraph_custom(
            doc, section["text"], bold=section.get("bold", False),
            italic=section.get("italic", False), size=section.get("size", 11),
            space_after=section.get("space_after", 6)
        )

    elif section_type == "bullet":
        add_bullet_list(doc, section["items"])

    elif section_type == "numbered":
        add_numbered_list(doc, section["items"])

    elif section_type == "table":
        add_table_from_data(doc, section["headers"], section["rows"])

    elif section_type == "code":
        add_code_block(doc, section["code"], section.get("caption", ""))

    # image_url / processed 类型已在 process_document_images 中处理，不单独渲染


def add_signature(doc):
    """在文档末尾添加记录人和日期。"""
    doc.add_paragraph()
    p = doc.add_paragraph("记录人：罗智峰，杨博文\n日期：2026年6月")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))


def create_document(title, sections, screenshots=None):
    """
    根据标题、章节配置和截图列表生成 Word 文档。
    """
    doc = Document()

    # 设置默认正文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei"
    )
    style.font.size = Pt(11)

    # 文档主标题
    main_title = doc.add_heading(title, level=0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in main_title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    sub = doc.add_paragraph("AI 驱动的产学研用一体化智慧协同系统")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()

    # 渲染正文
    for section in sections:
        render_section(doc, section)

    # 截图附录
    if screenshots:
        add_heading_custom(doc, "关键截图", level=1)
        for screenshot in screenshots:
            add_heading_custom(doc, screenshot["title"], level=2)
            add_screenshot(doc, screenshot["path"], screenshot.get("caption", ""))

    return doc


def generate_day_doc(config):
    """
    生成单日的 Word 文档。
    """
    content_file = os.path.join(OUTPUT_DIR, config["content_file"])
    if not os.path.exists(content_file):
        print(f"内容文件不存在：{content_file}")
        return

    with open(content_file, "r", encoding="utf-8") as f:
        content = f.read()

    # 解析 Markdown 内容
    sections = parse_markdown_content(content)

    # 处理文档内嵌图片（如下载 Day2 ER 图）
    doc_screenshots = process_document_images(sections, config["day"])

    # 合并手动配置的截图
    if config.get("screenshots"):
        doc_screenshots.extend(config["screenshots"])

    # 生成 Word 文档
    doc = create_document(config["title"], sections, doc_screenshots)
    add_signature(doc)

    output_path = os.path.join(OUTPUT_DIR, f"智汇桥项目开发日志_{config['day']}.docx")
    doc.save(output_path)
    print(f"已生成：{output_path}")


def main():
    """主入口：依次生成 Day1-Day5 的开发日志 Word 文档。"""
    for config in DAY_CONFIGS:
        print(f"\n正在生成 {config['day']} 文档...")
        generate_day_doc(config)


if __name__ == "__main__":
    main()
