#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据 Day 6 开发日志内容生成 Word 文档，并将对应截图嵌入到文档中。
图片与说明文字紧凑排列。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\day6_log.docx"

# 截图目录
IMAGE_DIR = r"C:\Users\luofeng\Desktop\实训\docs\day6"

# 截图配置：标题、文件名、说明
SECTIONS = [
    {
        "title": "4.1 登录页",
        "image": "day6-01-login-page.png",
        "desc": "登录页提供账号、密码输入与角色选择，登录成功后 Token 写入本地存储，并跳转至系统首页。"
    },
    {
        "title": "4.2 学生视角——科研项目列表页",
        "image": "day6-02-student-project-list.png",
        "desc": "列表页展示科研项目卡片，支持关键词搜索、项目类型筛选、分页浏览。学生可点击卡片进入项目详情。"
    },
    {
        "title": "4.3 学生视角——项目详情页",
        "image": "day6-03-student-project-detail.png",
        "desc": "项目详情页展示项目完整信息，包括项目简介、成员要求、预期成果、起止时间等。学生可点击「申请加入」。"
    },
    {
        "title": "4.4 学生视角——申请加入弹窗",
        "image": "day6-04-student-apply-dialog.png",
        "desc": "学生填写申请理由后提交，系统校验是否重复申请，并返回成功或失败提示。"
    },
    {
        "title": "4.5 学生视角——我的申请列表页",
        "image": "day6-05-student-my-applications.png",
        "desc": "我的申请页以表格形式展示所有申请记录，包含项目名、申请理由、状态标签、审核回复和申请时间。"
    },
    {
        "title": "4.6 学生视角——科研画像编辑页",
        "image": "day6-06-student-researcher-profile.png",
        "desc": "学生/教师可在此完善研究方向、专业技能、科研成果和个人简介，为后续智能推荐提供数据基础。"
    },
    {
        "title": "4.7 教师视角——发布科研项目页",
        "image": "day6-07-teacher-project-publish.png",
        "desc": "教师/管理员可发布新的科研项目，表单包含项目基本信息、成员要求、预期成果、计划成员数与起止时间。"
    },
    {
        "title": "4.8 教师视角——项目申请审核抽屉",
        "image": "day6-08-teacher-applications-drawer.png",
        "desc": "项目发布者在项目详情页点击「查看申请」后，弹出申请管理抽屉，可对每条申请执行「通过」或「拒绝」操作。"
    },
    {
        "title": "4.9 企业视角——发布企业需求页",
        "image": "day6-09-enterprise-demand-publish.png",
        "desc": "企业用户可发布技术攻关、成果转化、人才招聘、联合研发等类型的需求，填写需求描述、技术要求、预算范围与合作模式。"
    },
    {
        "title": "4.10 企业视角——企业需求列表页",
        "image": "day6-10-enterprise-demand-list.png",
        "desc": "企业需求列表页以卡片形式展示所有需求，支持按需求类型与行业领域筛选，点击查看详情。"
    }
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("6月22日项目开发日志", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("「智汇桥」——AI驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()  # 空行

    # 今日工作目标
    add_heading_custom(doc, "一、今日工作目标", level=1)
    goals = [
        "完成前端构建检查（npm run build），确保生产构建无类型错误。",
        "完成科研撮合模块前端页面开发：项目详情页申请审核、我的申请列表页、科研项目发布页、企业需求发布页、科研画像编辑页。",
        "修复科研项目发布接口 publisher_id 字段缺失导致的 500 错误。",
        "前后端联调验证注册、登录、项目列表、需求列表、项目申请、审核等核心流程。",
        "截取关键页面截图，整理到 docs/day6/ 目录，并写入开发日志。"
    ]
    for goal in goals:
        add_paragraph_custom(doc, f"• {goal}", space_after=4)

    # 今日完成内容概述
    add_heading_custom(doc, "二、今日完成内容", level=1)
    add_paragraph_custom(doc,
        "Day 6 完成了科研撮合模块前端剩余页面的开发与前后端联调，包括项目详情页申请审核、我的申请列表、"
        "科研项目发布、企业需求发布、科研画像编辑等页面，并修复了项目发布接口 publisher_id 缺失导致的 500 错误。")

    # 关键页面截图
    add_heading_custom(doc, "三、关键页面截图", level=1)
    add_paragraph_custom(doc,
        "以下截图为 Day 6 完成后，使用测试账号在本地开发环境实际访问各页面所得，图片统一存放于 docs/day6/ 目录。")

    for section in SECTIONS:
        add_heading_custom(doc, section["title"], level=2)

        image_path = os.path.join(IMAGE_DIR, section["image"])
        if os.path.exists(image_path):
            # 插入图片，宽度设置为 6 英寸以保持清晰
            doc.add_picture(image_path, width=Inches(6.0))
            # 设置图片居中对齐，并紧贴说明文字
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{section['image']}]", bold=True, color=RGBColor(0xFF, 0x00, 0x00))

        # 图片说明，紧跟图片
        add_paragraph_custom(doc, section["desc"], size=10, color=RGBColor(0x66, 0x66, 0x66),
                            italic=True, first_line_indent=0.2, space_after=12)

    # 遇到的问题与解决方法
    add_heading_custom(doc, "四、遇到的问题与解决方法", level=1)
    problems = [
        ("1", "后端启动时提示 Port 8081 was already in use", "本地已有历史后端进程占用 8081 端口", "直接使用已运行的后端服务，无需重复启动"),
        ("2", "科研项目发布接口返回 500", "research_project 表的 publisher_id 字段无默认值，前端未传递", "在 ResearchController.publishProject 中从 SecurityContext 获取当前登录用户 ID 并设置 publisherId"),
        ("3", "企业需求发布接口编译错误", "EnterpriseDemand 实体无 publisherId 字段", "删除对 setPublisherId 的调用，仅设置 enterpriseId"),
        ("4", "前端构建时部分 chunk 超过 500 kB", "Element Plus 等依赖整体打包", "当前不影响功能，后续可通过动态导入（import()）进一步优化")
    ]

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["序号", "问题描述", "原因分析", "解决方法"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)

    for num, desc, reason, solution in problems:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = desc
        row_cells[2].text = reason
        row_cells[3].text = solution
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)

    # 明日工作计划
    add_heading_custom(doc, "五、明日工作计划（2026年6月23日）", level=1)
    plans = [
        ("1", "继续资源流转模块前端页面联调（预约、审批、归还）", "资源预约全流程跑通"),
        ("2", "完善科研撮合模块细节（如项目状态自动更新、申请状态实时刷新）", "模块体验更稳定"),
        ("3", "开始首页数据看板与管理后台页面开发", "首页展示真实统计数据"),
        ("4", "编写接口测试用例或 Postman 集合", "便于后续回归测试")
    ]
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.style = "Light Grid Accent 1"
    plan_hdr = plan_table.rows[0].cells
    for i, h in enumerate(["序号", "工作内容", "预计产出"]):
        plan_hdr[i].text = h
        for p in plan_hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for num, work, output in plans:
        row = plan_table.add_row().cells
        row[0].text = num
        row[1].text = work
        row[2].text = output
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    # 备注
    add_heading_custom(doc, "六、备注", level=1)
    notes = [
        "今日已完成科研撮合模块前端全部核心页面开发与联调，学生申请、教师审核、项目发布、需求发布、科研画像等功能均验证通过。",
        "前端生产构建通过，无阻塞性错误。",
        "所有关键页面截图已保存至 docs/day6/ 目录，并嵌入本日志文档，便于后续汇报与复盘。",
        "资源流转模块后端已在前期完成，明日重点推进其前端联调与首页看板开发。"
    ]
    for note in notes:
        add_paragraph_custom(doc, f"• {note}", space_after=4)

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月22日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"Day 6 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
