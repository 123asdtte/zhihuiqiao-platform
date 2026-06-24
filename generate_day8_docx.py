#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据 Day 8 开发日志内容生成 Word 文档，并将对应截图嵌入到文档中。
图片与说明文字紧凑排列。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\day8_log.docx"

# 截图目录
IMAGE_DIR = r"C:\Users\luofeng\Desktop\实训\docs\day8"

# 截图配置：标题、文件名、说明
SECTIONS = [
    {
        "title": "4.1 首页数据看板",
        "image": "home_dashboard.png",
        "desc": "登录后进入系统首页，顶部欢迎横幅展示平台定位，下方统计卡片实时显示平台用户、科研项目、企业需求、闲置资源、学习资源、资源预约等核心指标，右侧快捷入口可快速跳转至各功能模块。"
    },
    {
        "title": "4.2 学习资源列表页",
        "image": "learning_resources.png",
        "desc": "通过侧边栏「教学辅助 → 学习资源」进入学习资源列表页。页面支持按资源类型、学科领域、难度等级和关键词进行筛选搜索，以卡片形式展示课程、视频、论文、图书、工具等学习资源，并显示封面、浏览量与「开始学习」入口。"
    },
    {
        "title": "4.3 学习中心页",
        "image": "learning_center.png",
        "desc": "学生点击「开始学习」后进入学习中心。页面顶部展示正在学习、已完成、我的收藏三项统计，下方标签页可切换查看不同状态的学习记录，支持更新学习进度、标记完成与删除记录等操作。"
    },
    {
        "title": "4.4 管理员数据看板",
        "image": "admin_dashboard.png",
        "desc": "管理员账号进入「系统管理 → 数据看板」后，可查看全平台运营数据。核心指标包括注册用户及角色分布、科研项目、企业需求、闲置资源；二级指标展示项目申请数、学习资源数、资源预约数；底部提供快捷管理入口。"
    }
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("6月23日项目开发日志", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("「智汇桥」——AI驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()  # 空行

    # 今日工作目标
    add_heading_custom(doc, "一、今日工作目标", level=1)
    goals = [
        "完成后端教学辅助模块：学习资源、学习记录、收藏相关实体、Mapper、Service、Controller。",
        "开发前端教学辅助页面：学习资源列表页与学习中心页。",
        "开发首页数据看板与管理后台数据看板页面，展示全平台统计数据。",
        "补充教学辅助与数据看板相关路由与侧边栏菜单配置。",
        "前后端联调验证教学辅助模块与数据看板功能。",
        "截取关键页面截图，整理到 docs/day8/ 目录，并生成 Word 开发日志。"
    ]
    for goal in goals:
        add_paragraph_custom(doc, f"• {goal}", space_after=4)

    # 今日完成内容概述
    add_heading_custom(doc, "二、今日完成内容", level=1)
    add_paragraph_custom(doc,
        "Day 8 完成了教学辅助模块的完整开发与数据看板页面建设。后端新增学习资源、学习记录、收藏相关接口，"
        "前端新增学习资源列表页与学习中心页，首页与管理员后台均接入数据看板统计接口，"
        "并通过本地开发环境完成前后端联调验证。")

    # 关键页面截图
    add_heading_custom(doc, "三、关键页面截图", level=1)
    add_paragraph_custom(doc,
        "以下截图为 Day 8 完成后，使用测试账号在本地开发环境实际访问各页面所得，图片统一存放于 docs/day8/ 目录。")

    for section in SECTIONS:
        add_heading_custom(doc, section["title"], level=2)

        image_path = os.path.join(IMAGE_DIR, section["image"])
        if os.path.exists(image_path):
            # 插入图片，宽度设置为 6 英寸以保持清晰
            doc.add_picture(image_path, width=Inches(6.0))
            # 设置图片居中对齐，并紧贴说明文字
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{section['image']}]", bold=True, color=RGBColor(0xFF, 0x00, 0x00))

        # 图片说明，紧跟图片
        add_paragraph_custom(doc, section["desc"], size=10, color=RGBColor(0x66, 0x66, 0x66),
                            italic=True, first_line_indent=0.2, space_after=12)

    # 遇到的问题与解决方法
    add_heading_custom(doc, "四、遇到的问题与解决方法", level=1)
    problems = [
        ("1", "后端启动报端口 8081 被占用", "之前已有一个后端进程在运行，未释放端口", "通过 netstat 找到占用进程并终止，重新启动后端服务"),
        ("2", "PowerShell 无法使用 < 重定向导入 SQL", "PowerShell 不支持 Unix 风格的输入重定向", "使用 Get-Content data.sql -Encoding UTF8 | mysql -u root -p123456 zhihuiqiao 管道方式导入"),
        ("3", "学习资源列表与学习记录首次为空", "数据库尚未初始化学习资源与学习记录数据", "在 data.sql 中补充 5 条学习资源模拟数据并手动导入，便于联调展示"),
        ("4", "首页存在未使用图标导入", "HomeView.vue 中导入了 Pointer 图标但未使用，可能导致构建警告", "移除 Pointer 图标导入，保持代码整洁")
    ]

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["序号", "问题描述", "原因分析", "解决方法"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)

    for num, desc, reason, solution in problems:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = desc
        row_cells[2].text = reason
        row_cells[3].text = solution
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)

    # 明日工作计划
    add_heading_custom(doc, "五、后续工作计划", level=1)
    plans = [
        ("1", "完善教学辅助模块详情页与收藏交互", "支持资源详情查看、收藏/取消收藏等完整功能"),
        ("2", "补充学习资源发布入口", "教师或管理员可在前端发布学习资源"),
        ("3", "开展全链路集成测试", "验证科研撮合、资源流转、教学辅助、数据看板四大模块的协同工作"),
        ("4", "准备项目总结与演示材料", "整理开发成果，输出项目总结文档")
    ]
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.style = "Light Grid Accent 1"
    plan_hdr = plan_table.rows[0].cells
    for i, h in enumerate(["序号", "工作内容", "预计产出"]):
        plan_hdr[i].text = h
        for p in plan_hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for num, work, output in plans:
        row = plan_table.add_row().cells
        row[0].text = num
        row[1].text = work
        row[2].text = output
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    # 备注
    add_heading_custom(doc, "六、备注", level=1)
    notes = [
        "今日已完成教学辅助模块后端（学习资源、学习记录、收藏）与前端（学习资源列表、学习中心）的开发。",
        "首页数据看板与管理员数据看板均成功接入 /api/dashboard/stats 接口，实时展示平台运营数据。",
        "所有关键页面截图已保存至 docs/day8/ 目录，并嵌入本日志文档。",
        "本地开发环境前后端联调通过，学习资源列表、开始学习、学习中心统计、数据看板等功能均验证正常。",
        "数据库中已补充学习资源模拟数据，便于后续功能演示与测试。"
    ]
    for note in notes:
        add_paragraph_custom(doc, f"• {note}", space_after=4)

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月23日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"Day 8 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
