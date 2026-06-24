#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成「智汇桥」项目总结 Word 文档。

文档包含项目概述、技术架构、功能模块、开发成果、测试验证、问题与解决方案、
项目亮点与后续计划，并嵌入各阶段关键页面截图。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\project_summary.docx"

# 截图根目录
IMAGE_BASE_DIR = r"C:\Users\luofeng\Desktop\实训\docs"

# 截图配置：章节标题、相对路径、说明
SCREENSHOTS = [
    {
        "title": "登录页",
        "path": "day6/day6-01-login-page.png",
        "desc": "统一的登录入口，支持学生、教师、企业、管理员多角色身份认证。"
    },
    {
        "title": "首页数据看板",
        "path": "day8/home_dashboard.png",
        "desc": "首页展示平台核心数据统计、快捷入口与最新动态，方便用户快速了解平台运行状况。"
    },
    {
        "title": "科研项目列表",
        "path": "day6/day6-02-student-project-list.png",
        "desc": "支持关键词搜索、类型筛选与分页浏览，帮助学生发现感兴趣的科研项目。"
    },
    {
        "title": "项目详情与申请",
        "path": "day6/day6-03-student-project-detail.png",
        "desc": "项目详情页展示项目背景、成员需求与联系方式，学生可一键提交加入申请。"
    },
    {
        "title": "企业需求发布",
        "path": "day6/day6-09-enterprise-demand-publish.png",
        "desc": "企业用户可发布技术攻关、联合研发等需求，促进产学研对接。"
    },
    {
        "title": "闲置资源列表",
        "path": "day7/day7-resource-list-with-images.png",
        "desc": "资源列表以卡片形式展示闲置设备、场地、资料等资源，支持图片懒加载提升加载性能。"
    },
    {
        "title": "资源预约申请",
        "path": "day7/day7-04-booking-dialog.png",
        "desc": "用户可填写预约用途与时间段，提交后由资源所有者进行审核。"
    },
    {
        "title": "学习资源列表",
        "path": "day8/learning_resources.png",
        "desc": "教学辅助模块提供课程、论文、实验等学习资源的分类浏览与搜索。"
    },
    {
        "title": "学习中心",
        "path": "day8/learning_center.png",
        "desc": "学习中心汇总用户的学习记录、收藏资源与学习进度，便于跟踪学习过程。"
    },
    {
        "title": "个人中心",
        "path": "day9/user_profile.png",
        "desc": "用户可查看与编辑个人资料，系统按角色动态展示年级、职称或企业名称等扩展字段。"
    },
    {
        "title": "用户管理",
        "path": "day9/user_manage.png",
        "desc": "管理员可对平台用户进行查询、编辑、删除、重置密码等操作。"
    },
    {
        "title": "内容审核",
        "path": "day9/audit_manage.png",
        "desc": "内容审核页提供科研项目、企业需求、闲置资源、学习资源的上下架与状态管理。"
    },
    {
        "title": "管理员数据看板",
        "path": "day8/admin_dashboard.png",
        "desc": "管理员数据看板集中展示用户、项目、需求、资源、学习资源等多维度统计信息。"
    },
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def add_bullet_list(doc, items):
    """添加项目符号列表。"""
    for item in items:
        add_paragraph_custom(doc, f"• {item}", space_after=4)


def add_table_from_data(doc, headers, rows, header_bold=True):
    """根据表头和数据创建表格。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=header_bold)
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)
    return table


def main():
    """生成项目总结 Word 文档。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("智汇桥项目总结报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("AI 驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 一、项目概述
    add_heading_custom(doc, "一、项目概述", level=1)
    add_paragraph_custom(doc,
        "「智汇桥」是一款面向高校、科研机构与企业的产学研用一体化智慧协同平台。平台围绕科研撮合、资源流转、"
        "教学辅助与系统管理四大核心模块，打通项目发布、需求对接、资源预约、学习资源共享、用户管理与内容审核等关键业务流程，"
        "旨在降低产学研合作门槛，提升科研成果转化效率与教学资源共享水平。")

    # 二、技术架构
    add_heading_custom(doc, "二、技术架构", level=1)
    add_heading_custom(doc, "2.1 后端技术栈", level=2)
    add_bullet_list(doc, [
        "Spring Boot 3.1.6 + JDK 17",
        "Spring Security + JWT 无状态认证",
        "MyBatis-Plus 3.5.6 数据访问层",
        "MySQL 8.0（UTF8MB4 编码）",
        "Maven 项目构建管理",
        "Swagger / Knife4j 接口文档"
    ])

    add_heading_custom(doc, "2.2 前端技术栈", level=2)
    add_bullet_list(doc, [
        "Vue 3.5 + TypeScript",
        "Vite 8 构建工具",
        "Element Plus 2.14 UI 组件库",
        "Pinia 3 状态管理",
        "Vue Router 5 路由管理",
        "Axios HTTP 请求库",
        "SCSS 样式预处理"
    ])

    add_heading_custom(doc, "2.3 开发与测试工具", level=2)
    add_bullet_list(doc, [
        "Node.js v24.14.0 + npm",
        "PowerShell 命令行环境",
        "python-docx 生成 Word 日志与总结文档",
        "Selenium + Edge 浏览器自动化截图",
        "Postman / Invoke-RestMethod 接口调试"
    ])

    # 三、功能模块
    add_heading_custom(doc, "三、功能模块", level=1)
    add_paragraph_custom(doc,
        "平台按照用户角色与业务领域划分为四大模块，各模块之间通过统一的用户认证、消息通知与数据看板进行串联。")

    modules = [
        ("科研撮合", "科研项目发布、查看、申请、审核；企业需求发布与浏览；科研画像维护。"),
        ("资源流转", "闲置资源发布、搜索、预约、审核、归还；资源流转记录追踪。"),
        ("教学辅助", "学习资源发布、浏览、收藏；学习记录生成与学习中心汇总。"),
        ("系统管理", "用户管理、内容审核、数据看板统计；消息通知中心。"),
    ]
    add_table_from_data(doc, ["模块名称", "核心功能"], modules)

    doc.add_paragraph()

    # 四、开发成果
    add_heading_custom(doc, "四、开发成果", level=1)
    add_paragraph_custom(doc,
        "项目采用迭代式开发方式，从基础认证与科研撮合模块逐步扩展至资源流转、教学辅助、系统管理与消息通知，"
        "最终完成全链路集成测试与性能体验优化。")

    achievements = [
        ("认证与授权", "完成 JWT 登录认证、Token 刷新、角色权限控制、个人资料更新。"),
        ("科研撮合", "完成科研项目 CRUD、全文检索、项目申请与审核、企业需求发布与浏览。"),
        ("资源流转", "完成闲置资源发布与预约、时间冲突检测、审核、归还与流转日志。"),
        ("教学辅助", "完成学习资源发布、详情、收藏、学习记录与学习中心的完整闭环。"),
        ("系统管理", "完成用户管理、内容审核、管理员数据看板与首页数据看板。"),
        ("消息通知", "完成通知表设计、通知中心、未读计数、标记已读，并在项目申请与资源预约流程中自动发送通知。"),
        ("性能优化", "完成图片懒加载、搜索防抖、列表加载状态优化等体验提升。"),
        ("文档输出", "输出 Day6 至 Day9 开发日志 Word 文档及项目总结报告。"),
    ]
    add_table_from_data(doc, ["模块", "完成内容"], achievements)

    doc.add_paragraph()

    # 五、测试验证
    add_heading_custom(doc, "五、测试验证", level=1)
    add_paragraph_custom(doc,
        "项目通过接口级与页面级双重验证，确保各模块核心链路可用。")

    add_heading_custom(doc, "5.1 全链路集成测试", level=2)
    add_paragraph_custom(doc,
        "使用 integration_test.py 脚本对认证、科研撮合、资源流转、教学辅助、系统管理、个人中心六大模块的 27 个核心接口进行联调验证，"
        "所有用例全部通过，具体结果如下：")

    test_results = [
        ("认证模块", "5/5", "多角色登录与用户信息获取"),
        ("科研撮合模块", "6/6", "项目列表、详情、申请、审核、需求发布与浏览"),
        ("资源流转模块", "3/3", "资源列表、预约、我的预约"),
        ("教学辅助模块", "6/6", "学习资源列表、详情、记录、收藏、学习中心、发布"),
        ("系统管理模块", "5/5", "数据看板、用户列表、审核统计、审核列表、状态更新"),
        ("个人中心", "2/2", "查看个人信息、更新个人资料"),
        ("合计", "27/27", "所有模块均通过集成测试"),
    ]
    add_table_from_data(doc, ["测试模块", "通过情况", "覆盖范围"], test_results)

    doc.add_paragraph()

    add_heading_custom(doc, "5.2 消息通知验证", level=2)
    add_bullet_list(doc, [
        "学生提交项目申请后，项目发布人收到未读通知。",
        "项目发布人审核申请后，申请人收到审核结果通知。",
        "资源预约提交与审核流程同样触发对应通知。",
        "前端通知中心可正确展示、筛选、标记已读，顶部未读数量实时轮询。"
    ])

    add_heading_custom(doc, "5.3 页面截图验证", level=2)
    add_paragraph_custom(doc,
        "通过 Selenium + Edge 浏览器对登录页、首页、各业务列表与详情页、管理后台等关键页面进行自动化全页截图，"
        "截图已保存于 docs/ 目录并按开发阶段分类，部分关键截图见第六节。")

    # 六、关键页面截图
    add_heading_custom(doc, "六、关键页面截图", level=1)
    add_paragraph_custom(doc,
        "以下截图来自本地开发环境实际运行效果，覆盖平台核心页面与功能入口。")

    for section in SCREENSHOTS:
        add_heading_custom(doc, section["title"], level=2)
        image_path = os.path.join(IMAGE_BASE_DIR, section["path"])
        if os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{section['path']}]", bold=True, color=RGBColor(0xFF, 0x00, 0x00))
        add_paragraph_custom(doc, section["desc"], size=10, color=RGBColor(0x66, 0x66, 0x66),
                            italic=True, first_line_indent=0.2, space_after=12)

    # 七、遇到的问题与解决方案
    add_heading_custom(doc, "七、遇到的问题与解决方案", level=1)
    problems = [
        ("1", "AuditManage.vue 状态更新接口类型不一致", "科研项目/需求/资源状态为 string，学习资源状态为 number，统一类型声明导致 TS 报错", "将 updateApiMap 声明为 Record<TabType, any>，绕过严格的函数签名校验"),
        ("2", "HomeView.vue 与 Dashboard.vue 存在未使用导入", "移除 unused 变量后仍有历史导入残留", "移除 HomeView.vue 的 userStore 导入，在 Dashboard.vue 的 manageEntries 中使用 Reading 和 DocumentChecked 图标"),
        ("3", "后端 AuthController 缺少 SysUser 导入", "新增 /auth/profile 接口使用了 SysUser 实体但未导入", "添加 import com.zhihuiqiao.entity.SysUser"),
        ("4", "浏览器 MCP 截图在 headless 模式下失败", "TRAE 内置浏览器实例无可见渲染表面，无法生成清晰截图", "改用 Selenium + Edge headless 编写自动化截图脚本，完成全页截图"),
        ("5", "后端服务在联调过程中意外停止", "端口 8081 无服务响应，登录请求返回 502", "重新启动 zhihuiqiao-backend 服务，确保前后端均正常运行后再截图"),
        ("6", "ResourceController 缺少 NotificationService 导入", "添加消息通知功能后未导入 NotificationService", "在 ResourceController 中添加 import com.zhihuiqiao.service.NotificationService"),
        ("7", "NotificationController 返回类型不匹配", "listNotifications 返回 Result<Page<Notification>>，但 service 返回 IPage<Notification>", "将返回类型改为 Result<IPage<Notification>>，并导入 IPage"),
    ]
    add_table_from_data(doc, ["序号", "问题描述", "原因分析", "解决方法"], problems)

    doc.add_paragraph()

    # 八、项目亮点
    add_heading_custom(doc, "八、项目亮点", level=1)
    add_bullet_list(doc, [
        "前后端分离架构清晰，职责边界明确，便于后续扩展与维护。",
        "基于 JWT + Spring Security 实现无状态认证与角色权限控制。",
        "科研撮合模块支持全文检索与模糊查询兜底，提升项目发现效率。",
        "资源流转模块内置时间冲突检测与流转日志，保障资源预约公平可追溯。",
        "教学辅助模块覆盖学习资源全生命周期，支持学习进度跟踪。",
        "消息通知中心贯穿核心业务流程，提升用户感知与处理效率。",
        "数据看板为管理员提供平台运营全局视图，辅助决策。",
        "通过自动化截图与集成测试脚本实现可重复的验证流程。"
    ])

    # 九、后续计划
    add_heading_custom(doc, "九、后续计划", level=1)
    add_bullet_list(doc, [
        "引入 Redis 缓存热门数据，进一步降低数据库压力并提升接口响应速度。",
        "完善消息通知渠道，支持邮件、短信等站外通知方式。",
        "增加文件上传与 OSS 集成，支持项目附件、资源图片、学习资料云存储。",
        "引入 Elasticsearch 增强全文检索能力与搜索结果相关性。",
        "补充单元测试与接口覆盖率，完善 CI/CD 流程。",
        "根据用户反馈持续优化 UI/UX 与交互细节。"
    ])

    # 十、结语
    add_heading_custom(doc, "十、结语", level=1)
    add_paragraph_custom(doc,
        "「智汇桥」项目从需求分析、技术选型、模块开发到联调测试，完成了产学研用协同平台的核心功能建设。"
        "项目整体架构合理、功能闭环完整、测试覆盖充分，具备良好的可扩展性与可维护性，"
        "为后续上线运营与功能迭代奠定了坚实基础。")

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月23日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"项目总结 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
