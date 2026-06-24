#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成「智汇桥」项目演示脚本 Word 文档。

文档包含演示目标、环境、账号、分场景操作流程、预期结果与对应截图，
用于项目结项答辩或成果展示。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\demo_script.docx"

# 截图根目录
IMAGE_BASE_DIR = r"C:\Users\luofeng\Desktop\实训\docs"

# 演示场景配置
SCENARIOS = [
    {
        "title": "场景一：学生浏览科研项目并提交加入申请",
        "steps": [
            ("打开浏览器，访问登录页，使用学生账号 student01 / 123456 登录。", "day6/day6-01-login-page.png"),
            ("进入首页，查看平台数据看板与快捷入口。", "day8/home_dashboard.png"),
            ("点击「科研撮合 → 科研项目」进入项目列表，浏览或搜索感兴趣的项目。", "day6/day6-02-student-project-list.png"),
            ("点击项目卡片进入项目详情页，查看项目介绍、需求与联系方式。", "day6/day6-03-student-project-detail.png"),
            ("点击「申请加入」，填写申请理由并提交。", "day6/day6-04-student-apply-dialog.png"),
            ("提交后可在「我的申请」中查看申请状态。", "day6/day6-05-student-my-applications.png"),
        ]
    },
    {
        "title": "场景二：教师审核项目申请并查看通知",
        "steps": [
            ("使用教师账号 teacher01 / 123456 登录系统。", "day6/day6-01-login-page.png"),
            ("点击顶部通知铃铛，进入消息通知中心，查看学生提交的申请通知。", "day9/user_profile.png"),
            ("进入「科研撮合 → 我的项目」，查看待处理申请列表。", "day6/day6-08-teacher-applications-drawer.png"),
            ("对申请进行审核：通过或拒绝，并可填写回复。", "day6/day6-08-teacher-applications-drawer.png"),
            ("审核结果将通过消息通知推送给学生。", ""),
        ]
    },
    {
        "title": "场景三：学生预约闲置资源",
        "steps": [
            ("使用学生账号登录，进入「资源流转 → 资源列表」。", "day7/day7-resource-list-with-images.png"),
            ("浏览卡片式资源列表，使用关键词或类型筛选目标资源。", "day7/day7-resource-list-with-images.png"),
            ("点击资源进入详情页，查看资源描述、可用时段与所有者信息。", "day7/day7-03-resource-detail-student.png"),
            ("点击「预约资源」，选择起止时间与用途后提交。", "day7/day7-04-booking-dialog.png"),
            ("系统提示预约提交成功，可在「我的预约」中跟踪状态。", "day7/day7-06-my-bookings.png"),
        ]
    },
    {
        "title": "场景四：教师审核资源预约",
        "steps": [
            ("使用教师账号登录，进入「资源流转 → 我的预约」或资源详情页。", "day7/day7-09-booking-management.png"),
            ("查看待审核的预约申请，确认时间段是否冲突。", "day7/day7-09-booking-management.png"),
            ("点击「通过」或「拒绝」，完成预约审核。", "day7/day7-09-booking-management.png"),
            ("审核通过后资源状态变为「已借出」，并记录流转日志。", "day7/day7-08-resource-detail-owner.png"),
        ]
    },
    {
        "title": "场景五：学生学习资源与收藏",
        "steps": [
            ("使用学生账号登录，进入「教学辅助 → 学习资源」。", "day8/learning_resources.png"),
            ("浏览学习资源列表，支持按类型与关键词搜索。", "day8/learning_resources.png"),
            ("点击资源进入详情页，可查看资源信息并开始学习。", "day9/learning_resource_detail.png"),
            ("点击「收藏」将资源加入个人收藏夹。", "day9/learning_resource_detail.png"),
            ("进入「学习中心」查看学习记录、收藏资源与学习进度。", "day8/learning_center.png"),
        ]
    },
    {
        "title": "场景六：管理员进行用户管理与内容审核",
        "steps": [
            ("使用管理员账号 admin / 123456 登录系统。", "day6/day6-01-login-page.png"),
            ("进入「系统管理 → 用户管理」，查询、编辑、重置密码或删除用户。", "day9/user_manage.png"),
            ("进入「系统管理 → 内容审核」，查看四类内容的审核统计。", "day9/audit_manage.png"),
            ("切换标签页，对科研项目、企业需求、闲置资源、学习资源进行上下架管理。", "day9/audit_manage.png"),
            ("进入「系统管理 → 数据看板」，查看平台多维度统计信息。", "day8/admin_dashboard.png"),
        ]
    },
    {
        "title": "场景七：消息通知中心",
        "steps": [
            ("任意角色登录后，顶部导航栏显示未读通知数量徽章。", "day9/user_profile.png"),
            ("点击铃铛图标进入消息通知中心，查看全部或仅未读通知。", "day9/user_profile.png"),
            ("点击单条通知可标记为已读，未读数量实时减少。", ""),
            ("点击「全部标记为已读」可一键清空未读。", ""),
        ]
    },
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def add_bullet_list(doc, items):
    """添加项目符号列表。"""
    for item in items:
        add_paragraph_custom(doc, f"• {item}", space_after=4)


def add_step(doc, index, description, image_path=None):
    """添加演示步骤，可选嵌入截图。"""
    add_paragraph_custom(doc, f"步骤 {index}：{description}", bold=False, size=11,
                         space_after=4, first_line_indent=0)
    if image_path:
        full_path = os.path.join(IMAGE_BASE_DIR, image_path)
        if os.path.exists(full_path):
            doc.add_picture(full_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
            add_paragraph_custom(doc, f"配图：{image_path}", size=9, color=RGBColor(0x66, 0x66, 0x66),
                                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{image_path}]", bold=True,
                                 color=RGBColor(0xFF, 0x00, 0x00), space_after=10)


def main():
    """生成项目演示脚本 Word 文档。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("智汇桥项目演示脚本", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("AI 驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 一、演示目标
    add_heading_custom(doc, "一、演示目标", level=1)
    add_paragraph_custom(doc,
        "本演示脚本用于展示「智汇桥」平台的核心业务流程与功能闭环，覆盖学生、教师、企业、管理员四类角色，"
        "重点呈现科研撮合、资源流转、教学辅助、系统管理与消息通知五大模块的协同工作效果。")

    # 二、演示环境
    add_heading_custom(doc, "二、演示环境", level=1)
    add_bullet_list(doc, [
        "后端服务：Spring Boot 应用运行于 http://localhost:8081",
        "前端服务：Vue 3 + Vite 开发服务器运行于 http://localhost:5175（或当前可用端口）",
        "数据库：MySQL 8.0，已初始化测试数据",
        "浏览器：Microsoft Edge（推荐）或 Google Chrome",
        "测试账号：student01 / teacher01 / enterprise01 / admin，密码均为 123456"
    ])

    # 三、演示账号
    add_heading_custom(doc, "三、演示账号", level=1)
    add_paragraph_custom(doc, "演示过程中根据场景切换以下账号：")
    accounts = [
        ("student01", "学生", "浏览项目、提交申请、预约资源、学习资源"),
        ("teacher01", "教师", "发布项目/资源、审核申请/预约、查看通知"),
        ("enterprise01", "企业", "发布企业需求"),
        ("admin", "管理员", "用户管理、内容审核、数据看板"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["账号", "角色", "演示权限"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for account, role, perm in accounts:
        row = table.add_row().cells
        row[0].text = account
        row[1].text = role
        row[2].text = perm
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    doc.add_paragraph()

    # 四、演示流程
    add_heading_custom(doc, "四、演示流程", level=1)
    add_paragraph_custom(doc,
        "以下场景按业务流程顺序组织，演示时可根据时间灵活调整。每个场景包含具体操作步骤与预期效果，"
        "并附带关键页面截图供参考。")

    for scenario_index, scenario in enumerate(SCENARIOS, start=1):
        add_heading_custom(doc, scenario["title"], level=2)
        add_paragraph_custom(doc, "预期目标：" + scenario["steps"][-1][0], italic=True,
                             color=RGBColor(0x66, 0x66, 0x66), space_after=8)
        for step_index, (desc, img) in enumerate(scenario["steps"], start=1):
            add_step(doc, step_index, desc, img)

    # 五、演示注意事项
    add_heading_custom(doc, "五、演示注意事项", level=1)
    add_bullet_list(doc, [
        "演示前请确认后端服务（端口 8081）与前端开发服务器已正常启动。",
        "同一浏览器登录多账号时建议使用隐私窗口，避免 Token 冲突。",
        "项目申请与资源预约具有幂等校验，重复操作可能提示「已提交过申请」或「该时间段已被预约」。",
        "若页面数据未刷新，可点击搜索或切换分页触发重新加载。",
        "消息通知的未读数量会每 30 秒自动轮询，也可手动刷新页面。"
    ])

    # 六、常见问题应对
    add_heading_custom(doc, "六、常见问题应对", level=1)
    faq = [
        ("登录后页面空白或跳转异常", "检查前端路由配置与本地开发服务器端口，清除浏览器缓存后重试。"),
        ("接口返回 401/403", "确认 Token 未过期，必要时重新登录；管理员接口需要 admin 角色。"),
        ("提交申请或预约失败", "查看错误提示，可能是重复申请、时间冲突或资源状态不可用。"),
        ("通知未实时刷新", "通知中心可手动下拉刷新，或等待 30 秒轮询间隔。"),
    ]
    faq_table = doc.add_table(rows=1, cols=2)
    faq_table.style = "Light Grid Accent 1"
    faq_hdr = faq_table.rows[0].cells
    for i, h in enumerate(["问题现象", "应对方法"]):
        faq_hdr[i].text = h
        for p in faq_hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for question, answer in faq:
        row = faq_table.add_row().cells
        row[0].text = question
        row[1].text = answer
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    doc.add_paragraph()

    # 七、结语
    add_heading_custom(doc, "七、结语", level=1)
    add_paragraph_custom(doc,
        "通过以上场景的演示，可以完整呈现「智汇桥」平台在产学研用协同中的核心价值："
        "学生可快速发现科研与学习资源，教师与企业可高效发布需求并管理申请预约，"
        "管理员可一站式管理平台用户与内容，消息通知机制确保各角色及时感知业务动态。")

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月23日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"演示脚本 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
