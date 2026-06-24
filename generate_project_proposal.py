#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成「智汇桥」项目方案 Word 文档。

数据来源：飞书文档 https://my.feishu.cn/wiki/QgGNw7Q7Piw0UykeV0xctFL5nnf
文档内容涵盖项目概述、主要功能、技术路线、数据库设计、小组分工、
实施规划、进度安排、预期成果与项目特色。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\智汇桥项目方案.docx"


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    elif level == 2:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    else:
        set_font(run, size=12, bold=True)
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0.25):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def add_bullet_list(doc, items):
    """添加项目符号列表。"""
    for item in items:
        add_paragraph_custom(doc, f"• {item}", space_after=4, first_line_indent=0.25)


def add_numbered_list(doc, items):
    """添加编号列表。"""
    for index, item in enumerate(items, start=1):
        add_paragraph_custom(doc, f"{index}. {item}", space_after=4, first_line_indent=0.25)


def add_table_from_data(doc, headers, rows, header_bold=True):
    """根据表头和数据创建表格。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=header_bold)
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)
    return table


def main():
    """生成项目方案 Word 文档。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("智汇桥——AI驱动的产学研用一体化智慧协同系统", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("项目方案文档")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 仓库地址
    add_paragraph_custom(doc, "仓库地址：", bold=True, space_after=2)
    add_paragraph_custom(doc, "https://github.com/123asdtte/zhihuiqiao-platform",
                         color=RGBColor(0x00, 0x66, 0xCC), italic=True,
                         space_after=12, first_line_indent=0)

    # 一、项目概述
    add_heading_custom(doc, "一、项目概述", level=1)

    add_heading_custom(doc, "1.1 项目名称", level=2)
    add_paragraph_custom(doc, "智汇桥——AI驱动的产学研用一体化智慧协同系统")

    add_heading_custom(doc, "1.2 项目目标", level=2)
    add_paragraph_custom(doc,
        "针对当前高校科研合作中存在的信息不对称、资源闲置浪费、学生实践机会少、校企对接难等问题，"
        "设计并实现一个连接教师、学生、实验室与校企资源的智慧协同平台。系统通过科研项目智能撮合、"
        "校园闲置资源流转、教学辅助个性化支持三大核心模块，为师-生-机协同提供一站式解决方案，"
        "提升科研合作效率、资源利用率与学生学习效果。")

    add_heading_custom(doc, "1.3 项目定位", level=2)
    add_paragraph_custom(doc,
        "本项目以“智能撮合 + 资源共享 + 个性化学习辅助”为核心，融合 Vue 前端、Spring Boot 后端、"
        "MyBatis-Plus 数据访问、大模型智能体接入等技术，构建一个面向高校师生与企业合作方的轻量化智慧校园协同平台。"
        "系统重点通过后端业务系统与智能体能力的集成，实现“业务系统负责数据管理，智能体负责学习辅助与推荐”的解耦设计，"
        "适合作为毕业设计、课程实训和项目答辩展示系统。")

    # 二、项目主要功能
    add_heading_custom(doc, "二、项目主要功能", level=1)

    add_heading_custom(doc, "2.1 用户与权限管理", level=2)
    add_paragraph_custom(doc,
        "支持学生、教师、企业、管理员等多种角色登录，完成用户信息维护、角色权限控制、账号状态管理等基础功能。")

    add_heading_custom(doc, "2.2 科研项目智能撮合", level=2)
    add_bullet_list(doc, [
        "教师/企业发布科研项目或技术需求",
        "学生基于科研画像匹配适合的项目",
        "支持项目申请、审核、成员管理全流程",
        "提供企业需求发布与检索功能"
    ])

    add_heading_custom(doc, "2.3 校园闲置资源流转", level=2)
    add_bullet_list(doc, [
        "师生发布实验设备、图书资料、办公用品等闲置资源",
        "支持资源预约、借用审批、归还记录",
        "自动生成资源流转日志"
    ])

    add_heading_custom(doc, "2.4 教学辅助个性化支持", level=2)
    add_bullet_list(doc, [
        "课程资源管理（课程、章节、知识点、学习资料）",
        "智能学习问答",
        "知识点讲解生成",
        "个性化学习计划",
        "练习题与错题反馈",
        "学习画像与效果评估"
    ])

    add_heading_custom(doc, "2.5 数据统计看板", level=2)
    add_paragraph_custom(doc,
        "通过图表展示科研项目数量、资源流转情况、用户活跃度、学习进度等核心数据，"
        "为管理员和教师提供决策支持。")

    # 三、技术路线
    add_heading_custom(doc, "三、技术路线", level=1)

    add_heading_custom(doc, "3.1 前端技术", level=2)
    add_paragraph_custom(doc, "Vue 3.4、Vite 5、Element Plus 2.7、Pinia、Vue Router、Axios、ECharts")

    add_heading_custom(doc, "3.2 后端技术", level=2)
    add_paragraph_custom(doc, "Spring Boot 3.1.6、JDK 17、MyBatis-Plus 3.5.6、Spring Security、JWT、Maven")

    add_heading_custom(doc, "3.3 数据库与存储", level=2)
    add_paragraph_custom(doc, "MySQL 8.0、Redis 7、MinIO（预留）")

    add_heading_custom(doc, "3.4 AI 智能体接入", level=2)
    add_paragraph_custom(doc, "扣子智能体平台、大模型 API、知识库、RAG 检索增强")

    add_heading_custom(doc, "3.5 工程化工具", level=2)
    add_paragraph_custom(doc, "Maven、Git、IDEA、Postman")

    # 四、数据库设计
    add_heading_custom(doc, "四、数据库设计", level=1)

    add_heading_custom(doc, "4.1 核心数据表", level=2)
    db_tables = [
        ("1", "sys_user", "用户表"),
        ("2", "sys_role", "角色表"),
        ("3", "sys_user_role", "用户角色关联表"),
        ("4", "researcher_profile", "科研画像表"),
        ("5", "research_project", "科研项目表"),
        ("6", "project_application", "项目申请表"),
        ("7", "enterprise_demand", "企业需求表"),
        ("8", "idle_resource", "闲置资源表"),
        ("9", "resource_booking", "资源预约表"),
        ("10", "resource_transfer_log", "资源流转记录表"),
        ("11", "learning_resource", "学习资源表"),
        ("12", "learning_record", "学习记录表"),
    ]
    add_table_from_data(doc, ["序号", "表名", "说明"], db_tables)

    doc.add_paragraph()

    # 五、小组分工
    add_heading_custom(doc, "五、小组分工", level=1)

    add_heading_custom(doc, "5.1 后端开发", level=2)
    add_numbered_list(doc, [
        "项目整体架构设计、数据库表结构设计、后端工程搭建；",
        "用户、课程、章节、知识点、学习记录等后端接口开发；",
        "MyBatis-Plus 实体类、Mapper、Service、Controller 编写；",
        "智能体接口接入，包括学习问答、学习计划、练习题生成等功能；",
        "数据库初始化脚本、接口测试、后端部署与运行调试；",
        "编写系统设计文档中数据库、后端架构、接口设计相关内容。"
    ])

    add_heading_custom(doc, "5.2 前端开发", level=2)
    add_numbered_list(doc, [
        "前端 Vue 项目初始化、页面路由设计、整体 UI 风格统一；",
        "学生端页面开发，包括学习问答、课程学习、学习计划、练习记录等页面；",
        "教师或管理员端页面开发，包括课程管理、资源管理、数据看板等页面；",
        "前后端接口联调、页面交互优化、Bug 修复；",
        "项目测试用例整理、操作手册编写、演示视频录制；",
        "答辩 PPT 制作、系统演示流程梳理和项目材料归档。"
    ])

    # 六、实施规划
    add_heading_custom(doc, "六、实施规划", level=1)

    add_heading_custom(doc, "第一阶段：需求分析与方案设计", level=2)
    add_numbered_list(doc, [
        "明确系统用户角色，包括学生、教师、企业、管理员；",
        "梳理核心功能，包括科研撮合、资源流转、教学辅助；",
        "完成系统总体架构设计、数据库设计、功能模块划分；",
        "确定智能体接入方式和技术方案。"
    ])

    add_heading_custom(doc, "第二阶段：基础环境搭建", level=2)
    add_numbered_list(doc, [
        "搭建 Spring Boot 后端工程；",
        "创建 Vue3 前端项目，配置路由、状态管理和 UI 组件库；",
        "初始化 MySQL 数据库，导入建表脚本和基础测试数据；",
        "配置 MyBatis-Plus、Redis、文件存储等基础环境。"
    ])

    add_heading_custom(doc, "第三阶段：核心功能开发", level=2)
    add_numbered_list(doc, [
        "完成用户管理、科研项目管理、资源管理、学习资源管理等基础接口；",
        "完成学习记录、学习计划、题库、错题反馈等业务模块；",
        "接入扣子智能体或大模型 API，实现智能问答和学习规划；",
        "开发前端主要页面，实现科研撮合、资源流转、学习辅助等功能；",
        "完成前后端接口联调。"
    ])

    add_heading_custom(doc, "第四阶段：系统联调与优化", level=2)
    add_numbered_list(doc, [
        "对系统进行完整流程测试，包括登录、项目申请、资源预约、智能问答；",
        "优化智能体提示词，提高回答准确性和学习建议可用性；",
        "优化页面交互体验，完善异常提示和空数据状态；",
        "编写测试报告、接口说明和系统设计文档。"
    ])

    add_heading_custom(doc, "第五阶段：总结与答辩准备", level=2)
    add_numbered_list(doc, [
        "整理项目源码、数据库脚本、部署说明和操作手册；",
        "录制系统演示视频；",
        "制作答辩 PPT；",
        "进行模拟演示，检查系统流程是否完整；",
        "完成最终项目归档与提交。"
    ])

    # 七、进度安排
    add_heading_custom(doc, "七、进度安排", level=1)

    add_heading_custom(doc, "第1周", level=2)
    add_numbered_list(doc, [
        "完成项目需求分析、系统方案设计和数据库设计；",
        "搭建后端项目和前端 Vue 基础项目；",
        "完成数据库建表脚本、初始化数据和基础实体类生成；",
        "初步完成用户、科研、资源等基础模块开发。"
    ])

    add_heading_custom(doc, "第2周", level=2)
    add_numbered_list(doc, [
        "完成科研撮合、资源流转、学习记录等核心业务接口；",
        "接入扣子智能体或大模型 API，完成智能问答基础功能；",
        "完成学生端主要页面和教师端管理页面开发；",
        "开始前后端联调，修复基础功能问题。"
    ])

    add_heading_custom(doc, "第3周", level=2)
    add_numbered_list(doc, [
        "完成学习画像、效果评估、数据统计看板等扩展功能；",
        "优化智能体提示词和系统交互体验；",
        "完成系统测试、Bug 修复和部署验证；",
        "整理项目文档、测试报告、操作手册和答辩 PPT。"
    ])

    # 八、预期成果
    add_heading_custom(doc, "八、预期成果", level=1)
    add_numbered_list(doc, [
        "一套可运行的产学研用一体化智慧协同 Web 系统；",
        "完整的前后端项目源码；",
        "MySQL 数据库建表脚本和初始化数据；",
        "智能体接入配置和调用日志管理功能；",
        "系统设计文档、数据库设计文档、接口说明文档；",
        "测试报告、操作手册、答辩 PPT 和项目演示视频。"
    ])

    # 九、项目特色
    add_heading_custom(doc, "九、项目特色", level=1)
    add_numbered_list(doc, [
        "产学研用一体化：打通科研项目、资源流转、教学辅助三大场景，形成完整闭环；",
        "智能体辅助学习：通过智能问答、知识点讲解和学习计划生成，提升学生自主学习效率；",
        "个性化学习路径：根据学生目标、学习记录和薄弱知识点，生成针对性学习计划；",
        "资源高效流转：盘活校园闲置资源，提高设备、图书等资产利用率；",
        "业务与 AI 解耦：后端系统负责用户、项目、资源等业务数据，智能体负责问答和生成能力，结构清晰；",
        "适合毕设展示：系统功能完整，覆盖数据库、后端、前端、智能体接入、数据可视化等多个技术点。"
    ])

    # 结语
    doc.add_paragraph()
    conclusion = doc.add_paragraph("文档整理：罗智峰\n日期：2026年6月23日")
    conclusion.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in conclusion.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"项目方案 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
