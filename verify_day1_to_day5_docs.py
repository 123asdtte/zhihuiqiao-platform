#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证 Day1-Day5 生成的 Word 文档是否完整、截图是否嵌入成功。
"""

import os
from docx import Document

OUTPUT_DIR = r"C:\Users\luofeng\Desktop\实训"

files = [
    "智汇桥项目开发日志_Day1.docx",
    "智汇桥项目开发日志_Day2.docx",
    "智汇桥项目开发日志_Day3.docx",
    "智汇桥项目开发日志_Day4.docx",
    "智汇桥项目开发日志_Day5.docx",
]

for filename in files:
    path = os.path.join(OUTPUT_DIR, filename)
    print(f"\n检查文件: {filename}")
    if not os.path.exists(path):
        print("  [错误] 文件不存在")
        continue

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    image_count = len(doc.inline_shapes)
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    print(f"  段落数: {paragraph_count}")
    print(f"  标题数: {heading_count}")
    print(f"  表格数: {table_count}")
    print(f"  图片数: {image_count}")

    if "[图片缺失" in text:
        print("  [警告] 文档中存在图片缺失提示")
    else:
        print("  [正常] 未发现图片缺失提示")

    if filename in ["智汇桥项目开发日志_Day2.docx", "智汇桥项目开发日志_Day5.docx"]:
        if image_count == 0:
            print(f"  [错误] {filename} 应包含截图但图片数为 0")
        else:
            print(f"  [正常] {filename} 截图已嵌入")
