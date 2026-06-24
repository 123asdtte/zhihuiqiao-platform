#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据 Day 7 开发日志内容生成 Word 文档，并将对应截图嵌入到文档中。
图片与说明文字紧凑排列。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\day7_log.docx"

# 截图目录
IMAGE_DIR = r"C:\Users\luofeng\Desktop\实训\docs\day7"

# 截图配置：标题、文件名、说明
SECTIONS = [
    {
        "title": "4.1 登录页",
        "image": "day7-01-login-page.png",
        "desc": "登录页使用统一认证布局，输入账号密码后自动写入 JWT Token 并跳转系统首页。"
    },
    {
        "title": "4.2 资源列表页",
        "image": "day7-02-resource-list.png",
        "desc": "资源列表页以卡片形式展示闲置资源，支持按资源类型、资源状态筛选和关键词搜索。侧边栏「资源流转」子菜单已新增「我的预约」和「发布闲置资源」入口。"
    },
    {
        "title": "4.3 资源详情页（学生视角）",
        "image": "day7-03-resource-detail-student.png",
        "desc": "资源详情页左侧展示资源大图与缩略图，右侧展示资源名称、类型、状态、价格、位置、描述、借用规则及「立即预约」按钮。"
    },
    {
        "title": "4.4 资源预约弹窗",
        "image": "day7-04-booking-dialog.png",
        "desc": "学生点击「立即预约」后弹出预约对话框，选择开始时间、结束时间并填写用途说明即可提交。"
    },
    {
        "title": "4.5 预约成功后的详情页",
        "image": "day7-05-booking-success.png",
        "desc": "预约提交成功后，页面下方的「预约记录」表格新增一条「待审批」记录，资源所有者可在该表格执行审批操作。"
    },
    {
        "title": "4.6 我的预约列表页",
        "image": "day7-06-my-bookings.png",
        "desc": "「我的预约」页以表格形式展示当前用户的所有预约记录，包含资源名称、预约时间段、用途、状态、审批回复及操作按钮。"
    },
    {
        "title": "4.7 发布闲置资源页",
        "image": "day7-07-resource-publish.png",
        "desc": "用户可在此发布新的闲置资源，填写资源名称、类型、位置、价格、图片、描述和借用规则，发布成功后跳转资源列表页。"
    },
    {
        "title": "4.8 资源详情页（所有者视角）",
        "image": "day7-08-resource-detail-owner.png",
        "desc": "教师/资源所有者进入自己发布的资源详情页，预约记录表格会显示审批操作入口（通过/拒绝/归还）。"
    },
    {
        "title": "4.9 预约审批管理",
        "image": "day7-09-booking-management.png",
        "desc": "资源所有者可在详情页对预约申请进行「通过」「拒绝」或「归还」操作，审批通过后资源状态自动变为「已借出」。"
    },
    {
        "title": "4.10 资源列表 AI 配图效果",
        "image": "day7-resource-list-with-images.png",
        "desc": "使用 TRAE text_to_image 接口为每条闲置资源生成对应 AI 图片，并写入数据库 images 字段。资源列表页卡片现在可以正常显示示波器、服务器、图书、投影仪、会议室等真实配图，不再显示占位图。同时清理了因 spring.sql.init.mode=always 导致的 70 余条重复资源记录，并将该配置改为 never，防止后续重启再次插入重复数据。"
    }
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("6月22日项目开发日志", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("「智汇桥」——AI驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()  # 空行

    # 今日工作目标
    add_heading_custom(doc, "一、今日工作目标", level=1)
    goals = [
        "完善资源列表页与资源详情页的交互体验。",
        "新增「我的预约」页面，展示当前用户的预约记录。",
        "新增「发布闲置资源」页面，支持用户发布可借用资源。",
        "在路由与侧边栏菜单中注册新页面入口。",
        "修复资源预约接口中的日期格式解析问题。",
        "完善后端「发布闲置资源」接口，自动从 SecurityContext 获取所有者 ID。",
        "完善后端「我的预约列表」接口，返回资源名称便于前端展示。",
        "前后端联调验证资源预约全流程。",
        "截取关键页面截图，整理到 docs/day7/ 目录，并写入开发日志。"
    ]
    for goal in goals:
        add_paragraph_custom(doc, f"• {goal}", space_after=4)

    # 今日完成内容概述
    add_heading_custom(doc, "二、今日完成内容", level=1)
    add_paragraph_custom(doc,
        "Day 7 完成了资源流转模块前端页面的开发与前后端联调，新增了「我的预约」和「发布闲置资源」页面，"
        "修复了资源预约日期格式解析问题，并完善了后端发布资源接口与我的预约列表接口。")

    # 关键页面截图
    add_heading_custom(doc, "三、关键页面截图", level=1)
    add_paragraph_custom(doc,
        "以下截图为 Day 7 完成后，使用测试账号在本地开发环境实际访问各页面所得，图片统一存放于 docs/day7/ 目录。")

    for section in SECTIONS:
        add_heading_custom(doc, section["title"], level=2)

        image_path = os.path.join(IMAGE_DIR, section["image"])
        if os.path.exists(image_path):
            # 插入图片，宽度设置为 6 英寸以保持清晰
            doc.add_picture(image_path, width=Inches(6.0))
            # 设置图片居中对齐，并紧贴说明文字
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{section['image']}]", bold=True, color=RGBColor(0xFF, 0x00, 0x00))

        # 图片说明，紧跟图片
        add_paragraph_custom(doc, section["desc"], size=10, color=RGBColor(0x66, 0x66, 0x66),
                            italic=True, first_line_indent=0.2, space_after=12)

    # 遇到的问题与解决方法
    add_heading_custom(doc, "四、遇到的问题与解决方法", level=1)
    problems = [
        ("1", "资源预约提交后后端返回「系统繁忙」", "前端 el-date-picker 返回 YYYY-MM-DD HH:mm:ss 格式字符串，后端 LocalDateTime 默认只支持 ISO 格式", "在 ResourceBooking 实体类的所有 LocalDateTime 字段上添加 @JsonFormat(pattern = \"yyyy-MM-dd HH:mm:ss\")"),
        ("2", "我的预约列表不显示资源名称", "后端 ResourceBooking 实体没有资源名称字段", "添加 @TableField(exist = false) 的 resourceName 字段，并在 listMyBookings 接口中填充"),
        ("3", "发布闲置资源接口未校验所有者", "后端直接保存前端传入的 ownerId，存在伪造风险", "在 publishResource 中从 SecurityContext 获取当前登录用户 ID 并强制设置为 ownerId"),
        ("4", "资源列表页全部显示占位图", "idle_resource 表积累大量重复记录，且 images 字段为无效占位 URL；spring.sql.init.mode=always 导致每次重启都插入重复数据", "使用 TRAE text_to_image 接口生成 AI 图片 URL 写入数据库；清理重复记录；将 spring.sql.init.mode 改为 never")
    ]

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["序号", "问题描述", "原因分析", "解决方法"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)

    for num, desc, reason, solution in problems:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = desc
        row_cells[2].text = reason
        row_cells[3].text = solution
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)

    # 明日工作计划
    add_heading_custom(doc, "五、明日工作计划（2026年6月23日）", level=1)
    plans = [
        ("1", "开发教学辅助模块后端（学习资源、学习记录、收藏）", "教学模块基础 API"),
        ("2", "开发教学辅助模块前端（学习资源列表、学习中心）", "教学模块页面与联调"),
        ("3", "首页数据看板与管理后台页面开发", "首页展示统计数据"),
        ("4", "编写接口测试与完善异常处理", "系统更稳定")
    ]
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.style = "Light Grid Accent 1"
    plan_hdr = plan_table.rows[0].cells
    for i, h in enumerate(["序号", "工作内容", "预计产出"]):
        plan_hdr[i].text = h
        for p in plan_hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for num, work, output in plans:
        row = plan_table.add_row().cells
        row[0].text = num
        row[1].text = work
        row[2].text = output
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    # 备注
    add_heading_custom(doc, "六、备注", level=1)
    notes = [
        "今日已完成资源流转模块前端全部核心页面开发与联调，资源列表、资源详情、预约申请、我的预约、发布资源、审批/归还等功能均验证通过。",
        "前端生产构建通过，无阻塞性错误。",
        "所有关键页面截图已保存至 docs/day7/ 目录，并嵌入本日志文档。",
        "补充了资源列表 AI 配图效果截图与问题记录。",
        "明日重点推进教学辅助模块开发与首页数据看板。"
    ]
    for note in notes:
        add_paragraph_custom(doc, f"• {note}", space_after=4)

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月22日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"Day 7 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
