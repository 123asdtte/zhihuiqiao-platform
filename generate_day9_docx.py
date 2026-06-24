#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据后续开发（Day 9）日志内容生成 Word 文档，并将对应截图嵌入到文档中。
图片与说明文字紧凑排列。
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 文档输出路径
OUTPUT_PATH = r"C:\Users\luofeng\Desktop\实训\day9_log.docx"

# 截图目录
IMAGE_DIR = r"C:\Users\luofeng\Desktop\实训\docs\day9"

# 截图配置：标题、文件名、说明
SECTIONS = [
    {
        "title": "4.1 学习资源详情页",
        "image": "learning_resource_detail.png",
        "desc": "学习资源详情页左侧展示资源封面大图，右侧展示资源名称、类型、学科、难度、浏览量、点赞数、描述与内容链接，并提供「开始学习」与「收藏」操作入口。"
    },
    {
        "title": "4.2 发布学习资源页",
        "image": "learning_resource_publish.png",
        "desc": "教师或管理员可通过「发布学习资源」页填写资源名称、类型、学科领域、难度等级、封面 URL、内容链接与描述，将优质学习资源发布到平台。"
    },
    {
        "title": "4.3 个人中心页",
        "image": "user_profile.png",
        "desc": "个人中心页左侧展示用户头像、用户名、角色与基础信息，右侧提供可编辑资料表单，支持按角色动态显示年级、职称或企业名称等扩展字段。"
    },
    {
        "title": "4.4 用户管理页",
        "image": "user_manage.png",
        "desc": "管理员进入「系统管理 → 用户管理」后，可按角色类型和关键词筛选查询平台用户，执行编辑、删除、重置密码等操作，并支持分页展示。"
    },
    {
        "title": "4.5 内容审核页",
        "image": "audit_manage.png",
        "desc": "内容审核页顶部展示科研项目、企业需求、闲置资源、学习资源四类内容数量统计，下方通过标签页分别管理各类内容的上下架与通过/关闭状态。"
    }
]


def set_font(run, font_name="Microsoft YaHei", size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    """统一设置字体格式。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    """添加自定义标题。"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 0:
        set_font(run, size=22, bold=True)
    elif level == 1:
        set_font(run, size=18, bold=True, color=RGBColor(0x20, 0x20, 0x20))
    else:
        set_font(run, size=14, bold=True, color=RGBColor(0x30, 0x33, 0x33))
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=RGBColor(0, 0, 0), italic=False,
                         align=None, space_after=6, first_line_indent=0):
    """添加自定义正文段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if align:
        p.alignment = align
    return p


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Microsoft YaHei")
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_heading("6月23日项目后续开发日志", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph("「智汇桥」——AI驱动的产学研用一体化智慧协同系统")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=14, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()  # 空行

    # 今日工作目标
    add_heading_custom(doc, "一、今日工作目标", level=1)
    goals = [
        "补全教学辅助模块剩余功能：学习资源详情页、学习资源发布页、收藏交互。",
        "完善系统管理模块：用户管理页、内容审核页。",
        "开发个人中心页，支持用户查看与编辑个人资料。",
        "补充后端管理接口：用户管理、内容审核、个人信息更新。",
        "更新前端路由与菜单配置，确保新增页面可正常访问。",
        "完成前后端联调测试，修复构建与运行过程中的问题。",
        "截取关键页面截图并生成 Word 开发日志。"
    ]
    for goal in goals:
        add_paragraph_custom(doc, f"• {goal}", space_after=4)

    # 今日完成内容概述
    add_heading_custom(doc, "二、今日完成内容", level=1)
    add_paragraph_custom(doc,
        "后续开发完成了系统管理模块与个人中心，并补全了教学辅助模块的详情与发布功能。"
        "后端新增 AdminController 提供用户管理与内容审核接口，AuthController 新增个人信息更新接口。"
        "前端新增/完善了用户管理、内容审核、个人中心、学习资源详情、学习资源发布等页面，"
        "并通过 Selenium + Edge 完成页面截图验证。")

    # 关键页面截图
    add_heading_custom(doc, "三、关键页面截图", level=1)
    add_paragraph_custom(doc,
        "以下截图为后续开发完成后，使用测试账号在本地开发环境实际访问各页面所得，图片统一存放于 docs/day9/ 目录。")

    for section in SECTIONS:
        add_heading_custom(doc, section["title"], level=2)

        image_path = os.path.join(IMAGE_DIR, section["image"])
        if os.path.exists(image_path):
            # 插入图片，宽度设置为 6 英寸以保持清晰
            doc.add_picture(image_path, width=Inches(6.0))
            # 设置图片居中对齐，并紧贴说明文字
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_custom(doc, f"[图片缺失：{section['image']}]", bold=True, color=RGBColor(0xFF, 0x00, 0x00))

        # 图片说明，紧跟图片
        add_paragraph_custom(doc, section["desc"], size=10, color=RGBColor(0x66, 0x66, 0x66),
                            italic=True, first_line_indent=0.2, space_after=12)

    # 遇到的问题与解决方法
    add_heading_custom(doc, "四、遇到的问题与解决方法", level=1)
    problems = [
        ("1", "AuditManage.vue 状态更新接口类型不一致", "科研项目/需求/资源状态为 string，学习资源状态为 number，统一类型声明导致 TS 报错", "将 updateApiMap 声明为 Record<TabType, any>，绕过严格的函数签名校验"),
        ("2", "HomeView.vue 与 Dashboard.vue 存在未使用导入", "移除 unused 变量后仍有历史导入残留", "移除 HomeView.vue 的 userStore 导入，在 Dashboard.vue 的 manageEntries 中使用 Reading 和 DocumentChecked 图标"),
        ("3", "后端 AuthController 缺少 SysUser 导入", "新增 /auth/profile 接口使用了 SysUser 实体但未导入", "添加 import com.zhihuiqiao.entity.SysUser"),
        ("4", "浏览器 MCP 截图在 headless 模式下失败", "TRAE 内置浏览器实例无可见渲染表面，无法生成清晰截图", "改用 Selenium + Edge headless 编写自动化截图脚本，完成全页截图"),
        ("5", "后端服务在联调过程中意外停止", "端口 8081 无服务响应，登录请求返回 502", "重新启动 zhihuiqiao-backend 服务，确保前后端均正常运行后再截图")
    ]

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["序号", "问题描述", "原因分析", "解决方法"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)

    for num, desc, reason, solution in problems:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = desc
        row_cells[2].text = reason
        row_cells[3].text = solution
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)

    # 后续工作计划
    add_heading_custom(doc, "五、后续工作计划", level=1)
    plans = [
        ("1", "全链路集成测试", "验证科研撮合、资源流转、教学辅助、系统管理四大模块的协同工作"),
        ("2", "性能与体验优化", "优化页面加载速度、图片懒加载、表单交互细节"),
        ("3", "完善消息通知", "新增系统通知、申请审核通知、预约状态变更通知"),
        ("4", "项目总结与演示准备", "整理开发成果，输出项目总结文档与演示脚本")
    ]
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.style = "Light Grid Accent 1"
    plan_hdr = plan_table.rows[0].cells
    for i, h in enumerate(["序号", "工作内容", "预计产出"]):
        plan_hdr[i].text = h
        for p in plan_hdr[i].paragraphs:
            for r in p.runs:
                set_font(r, size=10, bold=True)
    for num, work, output in plans:
        row = plan_table.add_row().cells
        row[0].text = num
        row[1].text = work
        row[2].text = output
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=9)

    # 备注
    add_heading_custom(doc, "六、备注", level=1)
    notes = [
        "今日已完成教学辅助模块剩余功能（详情、发布、收藏）与系统管理模块（用户管理、内容审核）的开发。",
        "个人中心页已开发完成，支持按角色动态显示扩展字段并同步更新用户信息。",
        "后端新增 AdminController 与 AuthController 接口，前端新增 admin.ts 与 auth.ts API 封装。",
        "所有新增页面已通过 npm run build 构建检查，并通过本地开发环境联调验证。",
        "关键页面截图已保存至 docs/day9/ 目录，并嵌入本日志文档。"
    ]
    for note in notes:
        add_paragraph_custom(doc, f"• {note}", space_after=4)

    # 记录人
    doc.add_paragraph()
    footer = doc.add_paragraph("记录人：罗智峰\n日期：2026年6月23日")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"Day 9 Word 文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
